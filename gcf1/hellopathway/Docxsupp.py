import docx.oxml
import os
import shutil
import datetime

def main(requests):

    requests_args = requests.args
    
    if requests_args and "new_assetname" in requests_args and "old_assetname" in requests_args and "error_description" in requests_args and "dlx_name" in requests_args:
        new_assetname = requests_args['new_assetname']
        old_assetname = requests_args['old_assetname']
        error_description = requests_args['error_description']
        dlx_name = requests_args['dlx_name']
        
    else:
        new_assetname = "error"
        old_assetname = "error"
        error_description = "error"
        dlx_name = "error"

    return f"So that's the Dokument {new_assetname}"

class Dokument_var_change():
    
    #----------------------------------------------------------------------------------------------
    # File Operations

    def file_operations(self,device, old_asst,new_asst,ticket_content):
        Current_Date = datetime.datetime.today().strftime ('%d-%b-%Y')
        src = 'C:/Users/barna/Desktop/dokwymxml/przykład'
        src_docx = 'C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx'
        
        dst = 'C:/Users/barna/Desktop/dokwymxml'
        dst1 = 'C:/Users/barna/Desktop/dokwymxml1/Dokument wymiany'

        os.rename(src=src_docx, dst=dst1 + " " + device + " - " + old_asst + ' na ' + new_asst + ' - ' + str(ticket_content) + ' - ' + str(Current_Date) +'.docx')
        shutil.copy(src=src + "/Dokument wymiany.docx",dst=dst)

    def var_change(self,dlx_name, dev_name, asset_new, asset_old, proces_new, proces_old, ip_new,
                   ip_old, sn_new, sn_old, location_new, error_desp, dlx_name_for_call, if_correct, if_repair, ticket_content):
        Current_Date = datetime.datetime.today().strftime ('%d-%m-%Y')

        document = docx.Document('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
        
#------------------------------------------- Exceptions --------------------------------------
        location_old = "Lager"
        
        if dev_name[0:2] == "COT" or dev_name[0:6] == "Printer" or dev_name == "COT P8010" or dev_name == "COT C3804" or dev_name == "COT T5308" or dev_name == "COT T5304" or dev_name == "COT T5000" or dev_name == "Printer" or dev_name == "CLP8301" or dev_name == "CLP 8301" or dev_name == "COT P8210" or dev_name == "COT T8308" or dev_name == "C-3408":
            ip_new = ip_old
            ip_old = " "
            
        elif dev_name[0:2] == "HDT" or dev_name[0:8] == "FIS CASIO" or dev_name[0:2] == "MFT":
            ip_new = ip_new
            ip_old = ip_old
            
        else:
            ip_new = ip_new
            ip_old = " "
            
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #print(paragraph.text)
                        
#-------------------------------------------- Old Variables -----------------------------------
                        # Nazwa Asset - stara
                        if 'AssetName.old' in paragraph.text:
                            paragraph.text = paragraph.text.replace("AssetName.old", asset_old)

                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Nazwa procesowa - stara
                        if 'Process_name.old' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Process_name.old", proces_old)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # IP - stare
                        if 'IP_address1.old' in paragraph.text:
                            paragraph.text = paragraph.text.replace("IP_address1.old", ip_old)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Serial Number - stary
                        if 'Serial_no.old' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Serial_no.old", sn_old)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Lokalizacja - stara
                        if 'Location.old' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Location.old", location_old)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                            
#-------------------------------------------- New Variables -----------------------------------
                        # Nazwa Asset - nowa
                        if 'AssetName.new' in paragraph.text:
                            paragraph.text = paragraph.text.replace("AssetName.new", asset_new)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Nazwa procesowa - nowa
                        if 'Process_name.new' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Process_name.new", proces_new)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # IP - nowy
                        if 'IP_address1.new' in paragraph.text:
                            paragraph.text = paragraph.text.replace("IP_address1.new", ip_new)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Serial Number - nowy
                        if 'Serial_no.new' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Serial_no.new", sn_new)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Lokalizacja - nowa
                        if 'Location.new' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Location.new", location_new)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                            
#-------------------------------------------- The Rest of Variables-----------------------------------
                        # Opis problemu
                        if 'error_description' in paragraph.text:
                            paragraph.text = paragraph.text.replace("error_description", error_desp)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Osoba wymieniająca
                        if 'dlx_name' in paragraph.text:
                            paragraph.text = paragraph.text.replace("dlx_name", dlx_name)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Osoba zgłaszająca
                        if 'applicant' in paragraph.text:
                            paragraph.text = paragraph.text.replace("applicant", dlx_name_for_call)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Rodzaj urządzenia
                        if 'Device' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Device", dev_name)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Czy wymiana przebiegła pomyślnie
                        if 'If_correct' in paragraph.text:
                            paragraph.text = paragraph.text.replace("If_correct", if_correct)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Czy urządzenie trzeba wysłac do naprawy
                        if 'If_repair' in paragraph.text:
                            paragraph.text = paragraph.text.replace("If_repair", if_repair)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Aktualna data
                        if 'Date' in paragraph.text:
                            paragraph.text = paragraph.text.replace("Date", Current_Date)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
                        # Ticket ID
                        if 'ticket_content' in paragraph.text:
                            paragraph.text = paragraph.text.replace("ticket_content", ticket_content)
                            print(paragraph.text)
                            document.save('C:/Users/barna/Desktop/dokwymxml/Dokument wymiany.docx')
